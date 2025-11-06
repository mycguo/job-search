"""
Interview Preparation Dashboard

Main hub for managing interview prep materials:
- Questions and answers
- Technical concepts
- Company research
- Practice sessions
"""

import streamlit as st
import sys
from datetime import datetime

# Add parent directory to path
sys.path.insert(0, '.')

from storage.interview_db import InterviewDB
from models.interview_prep import (
    create_interview_question,
    create_technical_concept,
    create_company_research
)
from simple_vector_store import SimpleVectorStore
from pages.app_admin import get_text_chunks
import io
from langchain_google_genai import ChatGoogleGenerativeAI
import json
from PyPDF2 import PdfReader
import docx


def extract_text_from_file(file_bytes, filename):
    """Extract text content from uploaded file (PDF, Word, or Text)"""
    try:
        file_lower = filename.lower()

        if file_lower.endswith('.pdf'):
            # Extract from PDF
            pdf = PdfReader(io.BytesIO(file_bytes))
            text = ""
            metadata = {
                'filename': filename,
                'num_pages': len(pdf.pages),
                'type': 'pdf'
            }

            for page in pdf.pages:
                text += page.extract_text()

            return text, metadata

        elif file_lower.endswith('.docx'):
            # Extract from Word document
            doc = docx.Document(io.BytesIO(file_bytes))
            text = ""
            metadata = {
                'filename': filename,
                'num_paragraphs': len(doc.paragraphs),
                'type': 'docx'
            }

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text, metadata

        elif file_lower.endswith('.txt'):
            # Extract from text file
            text = file_bytes.decode('utf-8')
            metadata = {
                'filename': filename,
                'type': 'txt'
            }

            return text, metadata

        else:
            return None, {'error': 'Unsupported file type'}

    except Exception as e:
        return None, {'error': str(e)}


def add_question_to_vector_store(question, answer, metadata):
    """Add question and answer to vector store for semantic search"""
    try:
        vector_store = SimpleVectorStore(store_path="./vector_store_personal_assistant")

        # Create searchable content
        content = f"""Interview Question: {question}

Answer: {answer}

Type: {metadata.get('type', '')}
Category: {metadata.get('category', '')}
Companies: {', '.join(metadata.get('companies', []))}
Tags: {', '.join(metadata.get('tags', []))}"""

        # Add to vector store
        text_chunks = get_text_chunks(content)
        metadatas = [{
            'source': 'interview_question',
            'question_id': metadata.get('question_id'),
            'type': 'interview_prep',
            'timestamp': datetime.now().isoformat(),
            **metadata
        } for _ in text_chunks]

        vector_store.add_texts(text_chunks, metadatas=metadatas)
        return True
    except Exception as e:
        st.error(f"Error adding to vector store: {str(e)}")
        return False


def add_document_to_vector_store(file_content, filename, metadata):
    """Add entire document to vector store for semantic search"""
    try:
        vector_store = SimpleVectorStore(store_path="./vector_store_personal_assistant")

        # Add document metadata
        content = f"""Interview Prep Document: {filename}

{file_content}

Category: {metadata.get('category', 'interview_prep')}
Tags: {', '.join(metadata.get('tags', []))}"""

        # Chunk and add to vector store
        text_chunks = get_text_chunks(content)
        metadatas = [{
            'source': 'interview_prep_document',
            'filename': filename,
            'type': 'interview_prep',
            'timestamp': datetime.now().isoformat(),
            **metadata
        } for _ in text_chunks]

        vector_store.add_texts(text_chunks, metadatas=metadatas)
        return True, len(text_chunks)
    except Exception as e:
        return False, str(e)


def parse_questions_from_document(text_content):
    """Use AI to parse questions and answers from document"""
    try:
        model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.0)

        prompt = f"""Parse this interview preparation document and extract all questions and answers.

Document:
{text_content}

Please extract all interview questions and their answers. Return a JSON array with this structure:
[
  {{
    "question": "the question text",
    "answer": "the answer text",
    "type": "behavioral or technical or system-design or case-study",
    "category": "best guess category (leadership, algorithms, conflict, etc.)"
  }}
]

Only return valid JSON, no other text. If no clear Q&A found, return empty array []."""

        response = model.invoke(prompt)

        # Parse JSON response
        content = response.content.strip()
        # Remove markdown code blocks if present
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]

        parsed = json.loads(content.strip())
        return True, parsed
    except Exception as e:
        return False, f"Error parsing document: {str(e)}"


def show_upload_document_form(db: InterviewDB):
    """Show form to upload interview prep document"""
    st.subheader("📄 Upload Interview Prep Document")

    st.markdown("""
    Upload a document containing your interview questions and answers.
    Supports: PDF, Word, Text files.
    """)

    with st.form("upload_document_form", clear_on_submit=True):
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['pdf', 'docx', 'txt'],
            help="Upload your interview prep notes, Q&A document, etc."
        )

        col1, col2 = st.columns(2)

        with col1:
            category = st.text_input(
                "Category (optional)",
                placeholder="e.g., behavioral, technical",
                help="Category for organization"
            )

        with col2:
            tags = st.text_input(
                "Tags (comma-separated, optional)",
                placeholder="e.g., leadership, algorithms",
                help="Tags for easier searching"
            )

        st.markdown("### Processing Options")

        process_mode = st.radio(
            "How do you want to process this document?",
            [
                "📦 Quick: Just store in searchable knowledge base (Recommended)",
                "🤖 Smart: AI parse into individual questions (slower)"
            ],
            help="Quick mode stores the whole document. Smart mode uses AI to extract Q&A pairs."
        )

        submit = st.form_submit_button("Upload Document", type="primary")

        if submit and uploaded_file:
            with st.spinner("Processing document..."):
                # Extract text from file
                file_bytes = uploaded_file.read()
                text_content, metadata = extract_text_from_file(file_bytes, uploaded_file.name)

                if not text_content:
                    st.error("Could not extract text from file")
                    return

                # Prepare metadata
                doc_metadata = {
                    'category': category if category else 'interview_prep',
                    'tags': [t.strip() for t in tags.split(',')] if tags else [],
                    'original_filename': uploaded_file.name
                }

                if "Quick" in process_mode:
                    # Quick mode: Just add to vector store
                    success, result = add_document_to_vector_store(
                        text_content,
                        uploaded_file.name,
                        doc_metadata
                    )

                    if success:
                        st.success(f"✅ Document uploaded! Added {result} chunks to knowledge base.")
                        st.info("💡 You can now ask questions about this document in the chat!")

                        with st.expander("📄 Document Preview"):
                            st.text(text_content[:1000] + "..." if len(text_content) > 1000 else text_content)
                    else:
                        st.error(f"❌ Error: {result}")

                else:
                    # Smart mode: Parse questions
                    st.info("🤖 Parsing questions with AI... This may take 10-20 seconds.")
                    success, result = parse_questions_from_document(text_content[:10000])  # Limit for API

                    if not success:
                        st.error(result)
                        st.info("Falling back to quick mode...")
                        success, chunks = add_document_to_vector_store(
                            text_content,
                            uploaded_file.name,
                            doc_metadata
                        )
                        if success:
                            st.success(f"✅ Document stored in knowledge base ({chunks} chunks)")
                        return

                    if not result or len(result) == 0:
                        st.warning("No questions found. Storing document as-is.")
                        success, chunks = add_document_to_vector_store(
                            text_content,
                            uploaded_file.name,
                            doc_metadata
                        )
                        if success:
                            st.success(f"✅ Document stored in knowledge base ({chunks} chunks)")
                        return

                    # Show parsed questions
                    st.success(f"✅ Found {len(result)} questions!")

                    with st.expander(f"📋 Review {len(result)} Parsed Questions", expanded=True):
                        questions_to_save = st.multiselect(
                            "Select questions to save:",
                            range(len(result)),
                            default=range(len(result)),
                            format_func=lambda i: f"Q{i+1}: {result[i]['question'][:60]}..."
                        )

                        for i in questions_to_save:
                            q_data = result[i]
                            st.markdown(f"**Q{i+1}:** {q_data['question']}")
                            st.caption(f"Type: {q_data.get('type', 'unknown')} | Category: {q_data.get('category', 'unknown')}")
                            with st.expander("View Answer"):
                                st.write(q_data['answer'])
                            st.divider()

                        if st.button("💾 Save Selected Questions", key="save_parsed"):
                            saved_count = 0
                            for i in questions_to_save:
                                q_data = result[i]

                                # Create question
                                q = create_interview_question(
                                    question=q_data['question'],
                                    type=q_data.get('type', 'technical'),
                                    category=q_data.get('category', category or 'general'),
                                    difficulty='medium',  # Default
                                    answer_full=q_data['answer'],
                                    tags=doc_metadata['tags']
                                )

                                # Save to database
                                q_id = db.add_question(q)

                                # Add to vector store
                                add_question_to_vector_store(
                                    question=q_data['question'],
                                    answer=q_data['answer'],
                                    metadata={
                                        'question_id': q_id,
                                        'type': q_data.get('type', 'technical'),
                                        'category': q_data.get('category', 'general'),
                                        'tags': doc_metadata['tags']
                                    }
                                )
                                saved_count += 1

                            st.success(f"✅ Saved {saved_count} questions!")
                            st.balloons()
                            st.rerun()


def show_add_question_form(db: InterviewDB):
    """Show form to add new interview question"""
    st.subheader("➕ Add New Question & Answer")

    with st.form("add_question_form", clear_on_submit=True):
        # Question details
        question = st.text_area(
            "Interview Question *",
            placeholder="e.g., Tell me about a time you led a difficult project",
            height=100,
            help="The interview question you want to prepare for"
        )

        col1, col2, col3 = st.columns(3)

        with col1:
            q_type = st.selectbox(
                "Type *",
                ["behavioral", "technical", "system-design", "case-study"],
                help="Type of interview question"
            )

        with col2:
            category = st.text_input(
                "Category *",
                placeholder="e.g., leadership, algorithms",
                help="Category for organization (leadership, conflict, algorithms, etc.)"
            )

        with col3:
            difficulty = st.selectbox(
                "Difficulty",
                ["easy", "medium", "hard"]
            )

        # Answer section
        st.markdown("### Your Answer")

        # Choose answer format
        answer_format = st.radio(
            "Answer Format",
            ["Full Answer", "STAR Format (Behavioral)"],
            help="STAR: Situation, Task, Action, Result"
        )

        if answer_format == "STAR Format (Behavioral)":
            situation = st.text_area(
                "Situation",
                placeholder="Describe the context and background...",
                height=80
            )
            task = st.text_area(
                "Task",
                placeholder="What was the challenge or goal?...",
                height=80
            )
            action = st.text_area(
                "Action",
                placeholder="What specific actions did you take?...",
                height=80
            )
            result = st.text_area(
                "Result",
                placeholder="What was the outcome? Include metrics if possible...",
                height=80
            )

            # Combine for full answer
            answer_full = f"""**Situation:** {situation}

**Task:** {task}

**Action:** {action}

**Result:** {result}"""

            answer_star = {
                'situation': situation,
                'task': task,
                'action': action,
                'result': result
            } if all([situation, task, action, result]) else None
        else:
            answer_full = st.text_area(
                "Your Answer *",
                placeholder="Write your complete answer here...",
                height=200
            )
            answer_star = None

        # Additional details
        col1, col2 = st.columns(2)

        with col1:
            companies = st.text_input(
                "Companies (comma-separated)",
                placeholder="e.g., Amazon, Google, Meta",
                help="Companies that ask this question"
            )

        with col2:
            tags = st.text_input(
                "Tags (comma-separated)",
                placeholder="e.g., leadership, team-management",
                help="Tags for easier searching"
            )

        notes = st.text_area(
            "Notes (optional)",
            placeholder="Additional notes, variations, tips...",
            height=80
        )

        confidence = st.slider(
            "Confidence Level",
            1, 5, 3,
            help="How confident are you in this answer? (1=Low, 5=High)"
        )

        # Vector store option
        add_to_vector = st.checkbox(
            "💾 Add to searchable knowledge base",
            value=True,
            help="Store in vector DB for semantic search and RAG queries"
        )

        submit = st.form_submit_button("Add Question", type="primary")

        if submit:
            if not question or not category or not answer_full:
                st.error("⚠️ Please fill in all required fields!")
            else:
                try:
                    # Parse companies and tags
                    company_list = [c.strip() for c in companies.split(',')] if companies else []
                    tag_list = [t.strip() for t in tags.split(',')] if tags else []

                    # Create question
                    q = create_interview_question(
                        question=question,
                        type=q_type,
                        category=category,
                        difficulty=difficulty,
                        answer_full=answer_full,
                        answer_star=answer_star,
                        notes=notes,
                        tags=tag_list,
                        companies=company_list,
                        confidence_level=confidence
                    )

                    # Save to database
                    q_id = db.add_question(q)
                    st.success(f"✅ Question added successfully! (ID: {q_id})")

                    # Add to vector store if requested
                    if add_to_vector:
                        with st.spinner("Adding to knowledge base..."):
                            success = add_question_to_vector_store(
                                question=question,
                                answer=answer_full,
                                metadata={
                                    'question_id': q_id,
                                    'type': q_type,
                                    'category': category,
                                    'companies': company_list,
                                    'tags': tag_list
                                }
                            )
                            if success:
                                st.success("✅ Added to searchable knowledge base!")

                    st.balloons()
                    st.rerun()

                except Exception as e:
                    st.error(f"Error: {str(e)}")


def show_recent_questions(db: InterviewDB, limit: int = 10):
    """Show recently added questions"""
    questions = db.list_questions()

    if not questions:
        st.info("No questions yet. Add your first one above!")
        return

    # Sort by created_at (most recent first)
    questions.sort(key=lambda x: x.created_at, reverse=True)
    questions = questions[:limit]

    st.subheader(f"📋 Recent Questions ({len(questions)})")

    for q in questions:
        with st.container():
            col1, col2, col3 = st.columns([4, 2, 1])

            with col1:
                st.markdown(f"**{q.question}**")
                # Badges
                badges = f"{q.get_display_type()} • {q.get_difficulty_emoji()} {q.difficulty.title()} • {q.category.title()}"
                if q.companies:
                    badges += f" • 🏢 {', '.join(q.companies[:2])}"
                st.caption(badges)

            with col2:
                st.write(f"Confidence: {q.get_confidence_emoji()} {q.confidence_level}/5")
                if q.practice_count > 0:
                    st.caption(f"Practiced {q.practice_count}x")

            with col3:
                if st.button("View", key=f"view_{q.id}"):
                    st.session_state['view_question_id'] = q.id
                    st.rerun()

            st.divider()


def main():
    st.set_page_config(page_title="Interview Prep", page_icon="🎯", layout="wide")

    st.title("🎯 Interview Preparation")
    st.markdown("Build your personal interview toolkit")

    # Initialize database
    db = InterviewDB()

    # Get stats
    stats = db.get_stats()

    # Key Metrics Row
    st.header("📊 Your Prep Stats")

    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        st.metric(
            "Questions",
            stats['total_questions'],
            help="Total questions in your bank"
        )

    with col2:
        st.metric(
            "Practiced",
            f"{stats['practice_percentage']:.0f}%",
            help="Percentage of questions you've practiced"
        )

    with col3:
        st.metric(
            "Concepts",
            stats['total_concepts'],
            help="Technical concepts stored"
        )

    with col4:
        st.metric(
            "Companies",
            stats['total_companies'],
            help="Company research entries"
        )

    with col5:
        st.metric(
            "Practice Hours",
            f"{stats['total_practice_time_hours']:.1f}",
            help="Total practice time"
        )

    st.divider()

    # Quick Actions
    st.header("⚡ Quick Actions")

    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        if st.button("📄 Upload Document", use_container_width=True):
            st.session_state['show_upload_document'] = True
            st.session_state['show_add_question'] = False

    with col2:
        if st.button("➕ Add Question", use_container_width=True):
            st.session_state['show_add_question'] = True
            st.session_state['show_upload_document'] = False

    with col3:
        if st.button("📝 View All Questions", use_container_width=True):
            # TODO: Navigate to questions page
            st.info("Questions page coming soon!")

    with col4:
        if st.button("💻 Technical Concepts", use_container_width=True):
            # TODO: Navigate to concepts page
            st.info("Concepts page coming soon!")

    with col5:
        if st.button("🎓 Practice Mode", use_container_width=True):
            # TODO: Navigate to practice page
            st.info("Practice mode coming soon!")

    st.divider()

    # Show upload document form if requested
    if st.session_state.get('show_upload_document', False):
        show_upload_document_form(db)

        if st.button("✕ Close Upload"):
            st.session_state['show_upload_document'] = False
            st.rerun()

        st.divider()

    # Show add question form if requested
    if st.session_state.get('show_add_question', False):
        show_add_question_form(db)

        if st.button("✕ Close Form"):
            st.session_state['show_add_question'] = False
            st.rerun()

        st.divider()

    # Breakdown by type
    if stats['total_questions'] > 0:
        st.header("📈 Question Breakdown")

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("By Type")
            for q_type, count in stats['questions_by_type'].items():
                percentage = (count / stats['total_questions'] * 100)
                st.write(f"**{q_type.title()}:** {count} ({percentage:.0f}%)")

        with col2:
            st.subheader("By Difficulty")
            for difficulty, count in stats['questions_by_difficulty'].items():
                percentage = (count / stats['total_questions'] * 100)
                emoji = {'easy': '🟢', 'medium': '🟡', 'hard': '🔴'}.get(difficulty, '⚪')
                st.write(f"**{emoji} {difficulty.title()}:** {count} ({percentage:.0f}%)")

        st.divider()

    # Recent questions
    show_recent_questions(db)

    # Navigation buttons
    st.divider()
    st.header("🧭 Navigation")

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("🏠 Home", use_container_width=True):
            st.switch_page("app.py")

    with col2:
        if st.button("📊 Dashboard", use_container_width=True):
            st.switch_page("pages/dashboard.py")

    with col3:
        if st.button("📝 Applications", use_container_width=True):
            st.switch_page("pages/applications.py")


if __name__ == "__main__":
    main()
